"""
Custom resume generator - creates tailored resumes for each job
Uses AI to optimize content and DOCX library to create professional documents
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
import os
import sys

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from ai.multi_provider_client import ai_client


class ResumeGenerator:
    def __init__(self):
        self.output_dir = "generated_resumes"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_custom_resume(
        self,
        resume_data: Dict,
        job_data: Dict,
        match_data: Dict
    ) -> str:
        """
        Generate a custom resume tailored to specific job
        
        Args:
            resume_data: User's resume information
            job_data: Target job information
            match_data: Match analysis results
        
        Returns: Path to generated DOCX file
        """
        # Use AI to rewrite resume sections
        optimized_summary = self._generate_summary(resume_data, job_data)
        optimized_experience = self._optimize_experience(resume_data, job_data)
        optimized_skills = self._prioritize_skills(resume_data, job_data, match_data)
        
        # Create DOCX
        doc = Document()
        
        # Header - Name and Contact
        self._add_header(doc, resume_data)
        
        # Professional Summary
        self._add_section(doc, "Professional Summary", optimized_summary)
        
        # Skills (prioritized for this job)
        self._add_section(doc, "Technical Skills", optimized_skills)
        
        # Experience (reordered and rewritten)
        self._add_section(doc, "Professional Experience", optimized_experience)
        
        # Education
        self._add_section(doc, "Education", resume_data.get('education', ''))
        
        # Save
        company = job_data.get('company', 'company').replace(' ', '_')
        job_id = job_data.get('id', 'job')
        filename = f"{self.output_dir}/resume_{company}_{job_id}.docx"
        doc.save(filename)
        
        return filename
    
    def _generate_summary(self, resume_data: Dict, job_data: Dict) -> str:
        """Generate tailored professional summary"""
        skills = resume_data.get('skills', [])
        if isinstance(skills, str):
            skills = [s.strip() for s in skills.split(',')]
        
        prompt = f"""
Write a professional summary (3-4 sentences) for a resume tailored to this job.
Highlight relevant skills and experience. Be specific and quantifiable.

Candidate Background:
- Skills: {', '.join(skills[:10])}
- Experience: {resume_data.get('experience_years', 0)} years
- Current Role: {resume_data.get('current_role', 'Professional')}

Target Job:
- Title: {job_data.get('title', '')}
- Company: {job_data.get('company', '')}
- Key Requirements: {job_data.get('requirements', '')[:200]}

Professional Summary:
"""
        
        try:
            return ai_client.chat_complete(
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200
            )
        except Exception as e:
            print(f"Error generating summary: {e}")
            return f"Experienced professional with {resume_data.get('experience_years', 0)} years in the field."
    
    def _optimize_experience(self, resume_data: Dict, job_data: Dict) -> str:
        """Rewrite experience section to highlight relevant achievements"""
        original_experience = resume_data.get('experience', '')
        
        prompt = f"""
Rewrite this work experience to emphasize skills relevant to the target job.
Use action verbs and quantify achievements. Keep it concise.

Original Experience:
{original_experience[:1000]}

Target Job: {job_data.get('title', '')}
Key Skills Needed: {job_data.get('requirements', '')[:200]}

Optimized Experience:
"""
        
        try:
            return ai_client.chat_complete(
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500
            )
        except Exception as e:
            print(f"Error optimizing experience: {e}")
            return original_experience
    
    def _prioritize_skills(self, resume_data: Dict, job_data: Dict, match_data: Dict) -> str:
        """Reorder skills to put matched skills first"""
        matched = match_data.get('matched_skills', [])
        all_skills = resume_data.get('skills', [])
        
        if isinstance(all_skills, str):
            all_skills = [s.strip() for s in all_skills.split(',')]
        
        # Put matched skills first
        prioritized = matched + [s for s in all_skills if s.lower() not in [m.lower() for m in matched]]
        
        # Group by category (optional enhancement)
        return ", ".join(prioritized[:15])  # Limit to top 15
    
    def _add_header(self, doc: Document, resume_data: Dict):
        """Add name and contact info header"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume_data.get('name', 'Your Name'))
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact
        contact_para = doc.add_paragraph()
        contact_text = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')} | {resume_data.get('location', '')}"
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(10)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
    
    def _add_section(self, doc: Document, title: str, content: str):
        """Add a section with title and content"""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
        
        # Horizontal line
        doc.add_paragraph('_' * 80)
        
        # Content
        content_para = doc.add_paragraph(content)
        content_para.style = 'Normal'
        
        doc.add_paragraph()  # Spacing


# Global instance
resume_generator = ResumeGenerator()
